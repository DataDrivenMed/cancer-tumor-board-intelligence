from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


@dataclass
class SourceSegment:
    segment_id: str
    text: str
    page: int | None = None
    paragraph: int | None = None


@dataclass
class ParsedDocument:
    document_id: str
    filename: str
    document_type: str
    segments: list[SourceSegment]

    @property
    def full_text(self) -> str:
        return "\n\n".join(segment.text for segment in self.segments)

    def numbered_text(self) -> str:
        rows: list[str] = []
        for segment in self.segments:
            locator = []
            if segment.page is not None:
                locator.append(f"page={segment.page}")
            if segment.paragraph is not None:
                locator.append(f"paragraph={segment.paragraph}")
            suffix = f"|{'|'.join(locator)}" if locator else ""
            rows.append(f"[{segment.segment_id}{suffix}] {segment.text}")
        return "\n".join(rows)


def _clean(text: str) -> str:
    return " ".join((text or "").split()).strip()


def parse_upload(filename: str, raw_bytes: bytes, document_id: str = "DOC-001") -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    segments: list[SourceSegment] = []

    if suffix == ".pdf":
        reader = PdfReader(BytesIO(raw_bytes))
        counter = 1
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            paragraphs = [p for p in page_text.splitlines() if _clean(p)]
            for paragraph_number, paragraph in enumerate(paragraphs, start=1):
                text = _clean(paragraph)
                if not text:
                    continue
                segments.append(
                    SourceSegment(
                        segment_id=f"S{counter:04d}",
                        text=text,
                        page=page_number,
                        paragraph=paragraph_number,
                    )
                )
                counter += 1
        return ParsedDocument(document_id, filename, "pdf", segments)

    if suffix == ".docx":
        doc = Document(BytesIO(raw_bytes))
        counter = 1
        for paragraph_number, paragraph in enumerate(doc.paragraphs, start=1):
            text = _clean(paragraph.text)
            if not text:
                continue
            segments.append(
                SourceSegment(
                    segment_id=f"S{counter:04d}",
                    text=text,
                    paragraph=paragraph_number,
                )
            )
            counter += 1
        return ParsedDocument(document_id, filename, "docx", segments)

    if suffix in {".txt", ".md"}:
        text = raw_bytes.decode("utf-8", errors="replace")
        counter = 1
        for paragraph_number, paragraph in enumerate(text.splitlines(), start=1):
            cleaned = _clean(paragraph)
            if not cleaned:
                continue
            segments.append(
                SourceSegment(
                    segment_id=f"S{counter:04d}",
                    text=cleaned,
                    paragraph=paragraph_number,
                )
            )
            counter += 1
        return ParsedDocument(document_id, filename, suffix.lstrip("."), segments)

    raise ValueError(f"Unsupported file type: {suffix}")


def parse_text(text: str, document_id: str = "DOC-001", filename: str = "pasted_case.txt") -> ParsedDocument:
    raw = text.encode("utf-8")
    return parse_upload(filename, raw, document_id=document_id)


def extract_text_from_upload(filename: str, raw_bytes: bytes) -> str:
    """Backward-compatible plain-text helper."""
    return parse_upload(filename, raw_bytes).full_text
